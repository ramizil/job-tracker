"""Parse an HTML resume into a matching profile (keywords + suggested queries).

The profile is a plain dict, persisted to data/profile.yaml so you can hand-edit
it (add synonyms, tune weights, add target titles) without touching code.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import yaml
from bs4 import BeautifulSoup

from . import config
from .config import PROFILE_PATH

# A curated skill dictionary. Keys are canonical skills; values are aliases that
# may appear in a job description. Extend freely in data/profile.yaml.
SKILL_ALIASES: dict[str, list[str]] = {
    "python": ["python", "py", "pytest"],
    "java": ["java", "testng", "spring boot", "spring"],
    "selenium": ["selenium"],
    "playwright": ["playwright"],
    "performance testing": ["performance", "load test", "loadrunner", "vugen", "jmeter", "nft", "non-functional"],
    "api testing": ["rest api", "rest", "api automation", "postman", "soapui"],
    "kafka": ["kafka", "msk"],
    "kubernetes": ["kubernetes", "k8s", "openshift", "aks", "eks", "helm"],
    "aws": ["aws", "amazon web services"],
    "azure": ["azure"],
    "docker": ["docker", "container"],
    "ci/cd": ["ci/cd", "cicd", "jenkins", "github actions", "gitlab ci", "argocd", "pipeline"],
    "observability": ["dynatrace", "grafana", "prometheus", "kibana", "observability", "monitoring"],
    "databases": ["oracle", "postgresql", "postgres", "mongodb", "couchbase", "sql", "n1ql"],
    "genai": ["genai", "llm", "gpt", "claude", "prompt engineering", "ai agent", "mcp", "cursor", "copilot"],
    "automation": ["automation", "qa automation", "sdet", "test automation"],
    "microservices": ["microservices", "microservice", "distributed systems"],
    "bash": ["bash", "shell", "ksh", "scripting"],
}

# Weight by category - skills you most want to be hired for score higher.
DEFAULT_WEIGHTS: dict[str, float] = {
    "performance testing": 3.0,
    "automation": 3.0,
    "genai": 2.5,
    "python": 2.5,
    "java": 2.0,
    "api testing": 2.0,
    "kubernetes": 1.5,
    "observability": 1.5,
    "microservices": 1.5,
}

# Job titles to target (used to build search queries).
DEFAULT_TARGET_TITLES = [
    "Performance Test Engineer",
    "Automation Engineer",
    "SDET",
    "QA Automation Engineer",
    "Performance Engineer",
    "AI / GenAI Engineer",
]


# Resume file types we can read text from.
SUPPORTED_RESUME_EXTS = {".html", ".htm", ".pdf", ".docx", ".txt", ".md", ".markdown"}


def _html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["style", "script"]):
        tag.decompose()
    return soup.get_text(separator=" ")


def _pdf_to_text(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "Reading PDF resumes needs PyMuPDF (pip install pymupdf)."
        ) from exc
    parts: list[str] = []
    with fitz.open(str(path)) as doc:
        for page in doc:
            parts.append(page.get_text())
    return "\n".join(parts)


def _docx_to_text(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "Reading Word resumes needs python-docx (pip install python-docx)."
        ) from exc
    d = docx.Document(str(path))
    parts: list[str] = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_text(path: Path | str) -> str:
    """Extract plain text from a resume file.

    Supports HTML, PDF, Word (.docx) and plain text/Markdown. Raises a clear
    error for missing files and unsupported types (e.g. legacy .doc).
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {path}")
    ext = path.suffix.lower()
    if ext in (".html", ".htm"):
        return _html_to_text(path.read_text(encoding="utf-8", errors="ignore"))
    if ext == ".pdf":
        return _pdf_to_text(path)
    if ext == ".docx":
        return _docx_to_text(path)
    if ext in (".txt", ".md", ".markdown", ""):
        return path.read_text(encoding="utf-8", errors="ignore")
    if ext == ".doc":
        raise ValueError(
            "Legacy .doc files aren't supported — open it in Word and 'Save As' "
            ".docx or PDF, then update the Resume path in Settings."
        )
    # Unknown extension: best-effort read as text.
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_text(path: Path) -> str:
    return extract_text(path).lower()


def build_profile(resume_path: Path | None = None) -> dict[str, Any]:
    """Parse the resume and return a profile dict (also saved to disk)."""
    path = Path(resume_path) if resume_path else config.RESUME_PATH
    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {path}")

    text = _read_text(path)

    detected: dict[str, list[str]] = {}
    for skill, aliases in SKILL_ALIASES.items():
        found = sorted({a for a in aliases if a in text})
        if found:
            detected[skill] = found

    profile: dict[str, Any] = {
        "resume_path": str(path),
        "skills": detected,
        "weights": {k: DEFAULT_WEIGHTS.get(k, 1.0) for k in detected},
        "target_titles": DEFAULT_TARGET_TITLES,
        "location": "Israel",
        "keywords": sorted(detected.keys()),
    }
    save_profile(profile)
    return profile


def save_profile(profile: dict[str, Any]) -> None:
    PROFILE_PATH.write_text(
        yaml.safe_dump(profile, sort_keys=False, allow_unicode=True),
        encoding="utf-8",
    )


def load_profile() -> dict[str, Any]:
    """Load the saved profile, building it from the resume on first use."""
    if PROFILE_PATH.exists():
        return yaml.safe_load(PROFILE_PATH.read_text(encoding="utf-8"))
    return build_profile()


def alias_map(profile: dict[str, Any]) -> dict[str, list[str]]:
    """Return {canonical_skill: [aliases]} limited to the profile's skills."""
    skills = profile.get("skills", {})
    # If hand-edited profile only lists keys, fall back to global aliases.
    out: dict[str, list[str]] = {}
    for skill in skills:
        out[skill] = SKILL_ALIASES.get(skill, [skill])
    return out


_word_re = re.compile(r"[a-z0-9][a-z0-9+#./-]*")


def tokenize(text: str) -> str:
    """Lowercased, whitespace-normalized text for substring matching."""
    return " ".join(_word_re.findall((text or "").lower()))
